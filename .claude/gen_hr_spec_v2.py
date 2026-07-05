# -*- coding: utf-8 -*-
"""PEP HR 员工主数据宽表 — 技术规格说明书 V2"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# ── 页面设置 ──
for sec in doc.sections:
    sec.page_width = Cm(29.7)
    sec.page_height = Cm(21.0)
    sec.left_margin = Cm(2.0)
    sec.right_margin = Cm(2.0)
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)

style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(9.5)
style.paragraph_format.space_after = Pt(3)
style.paragraph_format.line_spacing = 1.35

# ── 工具函数 ──
def h1(text):
    doc.add_heading(text, level=1)

def h2(text):
    doc.add_heading(text, level=2)

def h3(text):
    doc.add_heading(text, level=3)

def para(text):
    doc.add_paragraph(text)

def tbl(headers, rows):
    """统一风格的表格"""
    t = doc.add_table(rows=len(rows)+1, cols=len(headers), style='Table Grid')
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]
        c.text = h
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True; r.font.size = Pt(8.5)
    for i, rd in enumerate(rows):
        for j, v in enumerate(rd):
            c = t.rows[i+1].cells[j]
            c.text = str(v) if v is not None else ''
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8)
    return t

def code_block(sql):
    p = doc.add_paragraph()
    run = p.add_run(sql)
    run.font.name = 'Consolas'
    run.font.size = Pt(6.5)

# ================================================================
# 封面
# ================================================================
doc.add_paragraph(''); doc.add_paragraph(''); doc.add_paragraph('')
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('PEP 企业数据仓库'); r.font.size = Pt(30); r.bold = True
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('HR 员工主数据宽表'); r.font.size = Pt(22); r.bold = True
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('技术规格说明书'); r.font.size = Pt(18)
doc.add_paragraph('')
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('V1.0  |  2026-06-18  |  机密'); r.font.size = Pt(11)
r.font.color.rgb = RGBColor(128, 128, 128)
doc.add_paragraph(''); doc.add_paragraph('')
tbl(['项目', '内容'],
    [['项目名称', 'PEP 企业数据仓库 — HR 员工主数据宽表'],
     ['数据源', 'SAP HR 信息类型 → DataX → Hive dmp_ods'],
     ['目标范围', '公司代码 2000 | 集团 800 | 在职员工'],
     ['技术栈', 'Hive (Spark SQL) | dbt-core 1.11.9'],
     ['下游消费', 'FineReport BI 报表']])
doc.add_page_break()

# ================================================================
# 1. 文档概述
# ================================================================
h1('1. 文档概述')

h2('1.1 背景与目的')
para('PEP 企业数据仓库已通过 DataX 将 SAP HR 模块的员工信息类型（Infotype）同步至 Hive ODS 层。'
     '为支撑 FineReport BI 的人资报表需求，需将分散在 9 张表中的员工数据整合为一张宽表，'
     '按 pernr（人员编号）关联，统一口径，供下游直接消费。')
para('本文档定义该宽表的：数据源映射、字段字典、筛选逻辑、JOIN 策略、完整 SQL 及数据质量基线。')

h2('1.2 适用范围')
tbl(['维度', '取值', '备注'],
    [['集团', 'mandt = 800', 'SAP 标准集团'],
     ['公司代码', 'bukrs = 2000', '仅人民教育出版社本部'],
     ['时效性', 'begda < 当天 < endda', '当前有效记录（含 99991231 永不过期）'],
     ['员工状态', 'zhr_fjzt LIKE A%', '以 A 开头的在职/在岗状态'],
     ['时间基准', '2026-06-18', '本版报告采用的当前日期']])

h2('1.3 术语定义')
tbl(['术语', '全称', '说明'],
    [['Infotype', 'SAP HR 信息类型', 'SAP HR 模块以 4 位数字编码的数据组，如 PA0001=组织分配'],
     ['pernr', 'Personnel Number', 'SAP 员工唯一编号，跨信息类型的关联主键'],
     ['mandt', 'Client', 'SAP 集团编号，800 为 PEP 生产集团'],
     ['bukrs', 'Company Code', '公司代码，2000 为人民教育出版社'],
     ['begda/endda', '开始日期/结束日期', 'SAP HR 信息类型的有效期管理字段，yyyyMMdd 格式'],
     ['ZHR_*', '自定义增强字段', 'PEP 在 SAP 标准表上扩展的 Z 开头自定义字段']])

doc.add_page_break()

# ================================================================
# 2. 数据架构
# ================================================================
h1('2. 数据架构')

h2('2.1 数据流转')
para('SAP HR 模块（源）→ DataX 批量同步 → Hive dmp_ods（ODS 层，只读）→ 本宽表 SQL → DWD/DWS 层 → FineReport BI')
tbl(['阶段', '存储位置', '操作', '负责人'],
    [['源端', 'SAP ECC HR 模块', 'HR 业务人员日常录入', 'HR 部门'],
     ['同步', 'DataX 任务', '每日增量/全量同步至 Hive', '数据平台组'],
     ['ODS', 'Hive dmp_ods', '只读，不写入', '数据开发'],
     ['整合', '本宽表查询', '9 表 LEFT JOIN，按本文档规范执行', '数据开发'],
     ['消费', 'FineReport BI', '连接 Hive/OceanBase 展示', 'BI 开发']])

h2('2.2 源表全景')
tbl(['序号', 'SAP 信息类型', 'Hive 表名', '业务含义', '行数', '列数', '与主表关系', 'JOIN 覆盖率'],
    [['1', 'PA0001', 'ods_sap_pa0001', '组织分配', '12,318', '57', '锚表', '100%'],
     ['2', 'PA0002', 'ods_sap_pa0002', '个人数据', '5,286', '87', '1:1', '99.6%'],
     ['3', 'PA0105', 'ods_sap_pa0105', '通讯信息', '381,133', '26', '1:1（透转）', '99.5%'],
     ['4', 'PA0185', 'ods_sap_pa0185', '证件信息', '5,212', '43', '1:1', '99.4%'],
     ['5', 'PA0534', 'ods_sap_pa0534', '政治面貌', '3,351', '40', '1:1', '80.6%'],
     ['6', 'PA9101', 'ods_sap_pa9101', '教育经历（ZHR）', '2,782', '48', '1:N', '73.1%'],
     ['7', 'PA9112', 'ods_sap_pa9112', '资格认证（ZHR）', '767', '39', '1:N', '53.7%'],
     ['8', 'PA0023', 'ods_sap_pa0023', '工作经历', '2,799', '37', '1:N', '18.0%'],
     ['9', 'PA0021', 'ods_sap_pa0021', '家庭成员', '494', '80', '1:N', '14.7%']])

h2('2.3 实体关系图')
para('所有表通过 pernr（人员编号）与主表 PA0001 关联，采用 LEFT JOIN 策略以确保不丢主表数据：')
para('PA0001（锚）'
     ' ──LEFT── PA0002  [1:1]  '
     ' ──LEFT── PA0105  [1:1，透转]  '
     ' ──LEFT── PA0185  [1:1]  '
     ' ──LEFT── PA0534  [1:1]  '
     ' ──LEFT── PA9101  [1:N]  '
     ' ──LEFT── PA9112  [1:N]  '
     ' ──LEFT── PA0023  [1:N]  '
     ' ──LEFT── PA0021  [1:N]')

doc.add_page_break()

# ================================================================
# 3. 源表详细定义
# ================================================================
h1('3. 源表详细定义')

# 所有表共有的 SAP HR 标准头
h2('3.0 SAP HR 信息类型标准头（23 列，所有表共有）')
para('以下字段为 SAP HR 信息类型的元数据管理字段，每张 PA 表的前 23 列均为此结构。'
     '（注：ods_sap_pa0105 例外——列名全大写，数据类型为 varchar(512)）')
tbl(['序号', '列名', '中文名', '类型', '用途'],
    [['1', 'mandt', '集团', 'varchar(3)', '固定 800，所有表筛选条件'],
     ['2', 'pernr', '人员编号', 'varchar(8)', '关联主键，跨表 JOIN 唯一键'],
     ['3', 'subty', '子信息类型', 'varchar(4)', '如 MAIL/CELL/01，进一步细分记录类型'],
     ['4', 'objps', '对象标识', 'varchar(2)', ''],
     ['5', 'sprps', '锁定标识符', 'varchar(1)', ''],
     ['6', 'endda', '结束日期', 'varchar(8)', 'yyyyMMdd，99991231=永久有效'],
     ['7', 'begda', '开始日期', 'varchar(8)', 'yyyyMMdd，有效期起始'],
     ['8', 'seqnr', '记录号', 'varchar(3)', '同 pernr+begda+endda 下的序号'],
     ['9', 'aedtm', '更改日期', 'varchar(8)', '最后修改日期'],
     ['10', 'uname', '更改者', 'varchar(12)', '最后修改用户'],
     ['11', 'histo', '历史记录', 'varchar(1)', ''],
     ['12', 'itxex', '文本存在', 'varchar(1)', ''],
     ['13', 'refex', '参考存在', 'varchar(1)', ''],
     ['14', 'ordex', '确定字段存在', 'varchar(1)', ''],
     ['15', 'itbld', '屏幕控制', 'varchar(2)', ''],
     ['16', 'preas', '更改原因', 'varchar(2)', ''],
     ['17', 'flag1', '保留字段1', 'varchar(1)', ''],
     ['18', 'flag2', '保留字段2', 'varchar(1)', ''],
     ['19', 'flag3', '保留字段3', 'varchar(1)', ''],
     ['20', 'flag4', '保留字段4', 'varchar(1)', ''],
     ['21', 'rese1', '保留字段5', 'varchar(2)', ''],
     ['22', 'rese2', '保留字段6', 'varchar(2)', ''],
     ['23', 'grpvl', '分组值', 'varchar(4)', '']])

# ---- PA0001 ----
h2('3.1 PA0001 — 组织分配（锚表）')
para('记录员工的行政组织归属：所属公司、人事范围、组织单位、成本中心、岗位、职务等。'
     '本表为宽表的锚表（Anchor），所有筛选条件和 JOIN 均以本表为基准。')

h3('3.1.1 表级信息')
tbl(['项目', '内容'],
    [['Hive 表名', 'dmp_ods.ods_sap_pa0001'],
     ['全表行数', '12,318（mandt=800 下 10,405）'],
     ['过滤后基数', '851 人（bukrs=2000 + 当前有效 + A 状态）'],
     ['总列数', '57（= 23 标准头 + 34 业务列）'],
     ['复合主键', 'pernr + begda + endda + seqnr'],
     ['SAP 事务代码', 'PA30 / PA20'],
     ['增量键', 'begda']])

h3('3.1.2 业务字段')
tbl(['序号', '列名', '中文名', '类型', '填充率', '说明'],
    [['1', 'bukrs', '公司代码', 'varchar(4)', '100%', '筛选条件：2000'],
     ['2', 'werks', '人事范围', 'varchar(4)', '100%', '通常与公司代码一致'],
     ['3', 'btrtl', '人事子范围', 'varchar(4)', '100%', ''],
     ['4', 'persg', '员工组', 'varchar(1)', '100%', 'A=在职 D=兼职 E=外部'],
     ['5', 'persk', '员工子组', 'varchar(2)', '100%', '10/20/30/40/50/60'],
     ['6', 'orgeh', '组织单位', 'varchar(8)', '100%', '对应部门/科室'],
     ['7', 'plans', '职位', 'varchar(8)', '100%', '岗位编号'],
     ['8', 'stell', '职务', 'varchar(8)', '100%', '职务代码'],
     ['9', 'kostl', '成本中心', 'varchar(10)', '99.1%', '财务核算维度'],
     ['10', 'sname', '姓名', 'varchar(30)', '99.6%', '中文姓名'],
     ['11', 'ename', '英文名', 'varchar(40)', '99.6%', ''],
     ['12', 'zhr_ygzt', '员工状态', 'varchar(1)', '100%', 'A=在职 B=离职'],
     ['13', 'zhr_fjzt', '附加状态', 'varchar(3)', '100%', '筛选条件 A 开头'],
     ['14', 'zhr_zylx', '专业类型', 'varchar(2)', '100%', '02/03 等'],
     ['15', 'zhr_edlst', '最高学历', 'varchar(2)', '61.5%', '字典值'],
     ['16', 'zhr_dglst', '最高学位', 'varchar(1)', '61.3%', '字典值'],
     ['17', 'vdsk1', '组织代码', 'varchar(14)', '100%', '层次化编码'],
     ['18', 'gsber', '业务范围', 'varchar(4)', '100%', ''],
     ['19', 'abkrs', '工资范围', 'varchar(2)', '100%', ''],
     ['20', 'ansvh', '工作合同', 'varchar(2)', '0%', '全空，SAP 未启用'],
     ['21', 'mstbr', '主管', 'varchar(8)', '100%', ''],
     ['22', 'kokrs', '成本控制范围', 'varchar(4)', '100%', ''],
     ['23', 'otype', '对象类型', 'varchar(2)', '100%', 'S=员工'],
     ['24', 'sbmod', '组', 'varchar(4)', '100%', ''],
     ['25', 'fistl', '资金中心', 'varchar(16)', '100%', ''],
     ['26', 'geber', '基金', 'varchar(10)', '100%', ''],
     ['27', 'fkber', '功能范围', 'varchar(16)', '100%', ''],
     ['28', 'grant_nbr', '拨款编号', 'varchar(20)', '100%', ''],
     ['29', 'sgmnt', '段', 'varchar(10)', '100%', ''],
     ['30', 'budget_pd', '预算期间', 'varchar(10)', '100%', ''],
     ['31', 'juper', '法律', 'varchar(4)', '100%', ''],
     ['32', 'sacha', '工资核算管理员', 'varchar(3)', '100%', ''],
     ['33', 'sachp', '人事管理员', 'varchar(3)', '100%', ''],
     ['34', 'sachz', '时间管理员', 'varchar(3)', '100%', '']])

# ---- PA0002 ----
h2('3.2 PA0002 — 个人数据')
para('记录员工的个人基础信息：姓名、性别、出生信息、国籍、民族、籍贯、婚姻、宗教信仰、'
     '健康状况、研究方向、外语能力等。通过 pernr 与 PA0001 做 1:1 LEFT JOIN。')

h3('3.2.1 表级信息')
tbl(['项目', '内容'],
    [['Hive 表名', 'dmp_ods.ods_sap_pa0002'],
     ['全表行数', '5,286'],
     ['JOIN 覆盖率', '848 / 851 = 99.6%'],
     ['总列数', '87（= 23 标准头 + 64 业务列）']])

h3('3.2.2 业务字段')
tbl(['序号', '列名', '中文名', '类型', '填充率', '说明'],
    [['1', 'nachn', '姓', 'varchar(40)', '99.6%', ''],
     ['2', 'vorna', '名', 'varchar(40)', '73.1%', '179人缺名'],
     ['3', 'cname', '中文姓名', 'varchar(80)', '—', ''],
     ['4', 'gesch', '性别', 'varchar(1)', '99.6%', '1=男 2=女'],
     ['5', 'gbdat', '出生日期', 'varchar(8)', '99.6%', 'yyyyMMdd'],
     ['6', 'gbort', '出生地', 'varchar(40)', '—', ''],
     ['7', 'gblnd', '出生国', 'varchar(3)', '—', ''],
     ['8', 'natio', '国籍', 'varchar(3)', '99.6%', ''],
     ['9', 'famst', '婚姻状况', 'varchar(1)', '19.2%', '1=未婚 2=已婚 80%为空'],
     ['10', 'famdt', '结婚日期', 'varchar(8)', '—', ''],
     ['11', 'anzkd', '子女数', 'decimal(3,0)', '99.6%', ''],
     ['12', 'konfe', '宗教', 'varchar(2)', '—', ''],
     ['13', 'sprsl', '语言', 'varchar(1)', '—', ''],
     ['14', 'anred', '称谓', 'varchar(1)', '—', ''],
     ['15', 'perid', '身份证号', 'varchar(20)', '—', ''],
     ['16', 'titel', '头衔', 'varchar(15)', '—', ''],
     ['17', 'inits', '首字母', 'varchar(10)', '—', ''],
     ['18', 'rufnm', '称呼名', 'varchar(40)', '—', ''],
     ['19', 'zhr_mz', '民族', 'varchar(2)', '99.6%', '00=汉族 56民族+2'],
     ['20', 'zhr_hkxz', '户口性质', 'varchar(1)', '99.6%', ''],
     ['21', 'zhr_jgss', '籍贯省', 'varchar(8)', '99.6%', ''],
     ['22', 'zhr_jgs', '籍贯市', 'varchar(8)', '99.6%', ''],
     ['23', 'zhr_jgq', '籍贯区', 'varchar(8)', '—', ''],
     ['24', 'zhr_csss', '出生省', 'varchar(8)', '—', ''],
     ['25', 'zhr_css', '出生市', 'varchar(8)', '—', ''],
     ['26', 'zhr_csq', '出生区', 'varchar(8)', '—', ''],
     ['27', 'zhr_sthea', '健康状况', 'varchar(40)', '0%', '全空'],
     ['28', 'zhr_relig', '宗教信仰', 'varchar(40)', '0%', '全空'],
     ['29', 'zhr_speci', '研究方向或专长', 'varchar(40)', '—', ''],
     ['30', 'zhr_forla', '外语能力', 'varchar(60)', '—', ''],
     ['31', 'zhr_wybrk', '外语说明', 'varchar(255)', '—', ''],
     ['32', 'zhr_tssf', '特殊身份', 'varchar(1)', '—', ''],
     ['33', 'zhr_dasf', '档案身份', 'varchar(1)', '—', ''],
     ['34', 'zhr_heduc', '最高学历', 'varchar(2)', '0.8%', '主数据在PA0001'],
     ['35', 'zhr_hdegr', '最高学位', 'varchar(1)', '0.8%', '主数据在PA0001']])

doc.add_page_break()

# ---- PA0105 ----
h2('3.3 PA0105 — 通讯信息')
para('记录员工的通讯方式（邮箱、手机、座机等）。此表为行式存储，'
     '通过 subty 列区分通讯类型，需在查询时做行转列（PIVOT）处理。')
para('⚠️ 此表列名全大写且数据类型为 varchar(512)，与其他表风格不一致，为 DataX 早期同步配置差异所致。')

h3('3.3.1 表级信息')
tbl(['项目', '内容'],
    [['Hive 表名', 'dmp_ods.ods_sap_pa0105'],
     ['全表行数', '381,133'],
     ['JOIN 覆盖率', '847 / 851 = 99.5%'],
     ['总列数', '26（= 23 标准头 + 3 业务列）'],
     ['PIVOT 策略', "GROUP BY pernr + MAX(CASE WHEN subty='X' THEN usrid END)"]])
h3('3.3.2 subty 枚举及覆盖率')
tbl(['subty', '含义', '输出别名', '覆盖率（847人）', '是否输出'],
    [['MAIL', '电子邮箱', '邮箱', '841人 (98.8%)', '✅'],
     ['CELL', '手机号码', '手机号', '692人 (81.3%)', '✅'],
     ['MPHN', '座机号码', '座机号', '847人 (99.5%)', '✅'],
     ['CARD', '工卡号', '-', '833人 (97.9%)', '否'],
     ['0001', '系统账号', '-', '805人 (94.6%)', '否'],
     ['9001', '其他标识', '-', '477人 (56.1%)', '否']])

# ---- PA0185 ----
h2('3.4 PA0185 — 证件信息')
para('记录员工的各类证件信息。查询时仅取 subty=01（身份证），其余证件类型暂不纳入宽表。')

h3('3.4.1 表级信息')
tbl(['项目', '内容'],
    [['Hive 表名', 'dmp_ods.ods_sap_pa0185'],
     ['全表行数', '5,212'],
     ['JOIN 覆盖率', '846 / 851 = 99.4%'],
     ['总列数', '43（= 23 标准头 + 20 业务列）'],
     ['筛选条件', 'subty = 01（仅身份证）']])
h3('3.4.2 业务字段')
tbl(['序号', '列名', '中文名', '类型', '说明'],
    [['1', 'subty', '证件类型', 'varchar(4)', '01=身份证，筛选条件'],
     ['2', 'icnum', '证件号码', 'varchar(30)', ''],
     ['3', 'auth1', '签发机构', 'varchar(30)', ''],
     ['4', 'isspl', '发证机关', 'varchar(30)', ''],
     ['5', 'usefr', '有效期开始', 'varchar(8)', 'yyyyMMdd'],
     ['6', 'useto', '有效期截至', 'varchar(8)', 'yyyyMMdd']])

# ---- PA0534 ----
h2('3.5 PA0534 — 政治面貌')
para('记录员工的政治面貌、入党/入团信息。')

h3('3.5.1 表级信息')
tbl(['项目', '内容'],
    [['Hive 表名', 'dmp_ods.ods_sap_pa0534'],
     ['全表行数', '3,351'],
     ['JOIN 覆盖率', '686 / 851 = 80.6%'],
     ['总列数', '40（= 23 标准头 + 17 业务列）']])
h3('3.5.2 业务字段')
tbl(['序号', '列名', '中文名', '类型', '说明'],
    [['1', 'pcode', '政治面貌代码', 'varchar(2)', ''],
     ['2', 'joinu', '加入组织名称', 'varchar(40)', ''],
     ['3', 'intr1', '介绍人1', 'varchar(8)', ''],
     ['4', 'intr2', '介绍人2', 'varchar(8)', ''],
     ['5', 'paety', '党员类型', 'varchar(2)', ''],
     ['6', 'pstat', '状态', 'varchar(1)', ''],
     ['7', 'joipd', '加入日期', 'varchar(8)', 'yyyyMMdd'],
     ['8', 'bfmdt', '转正日期', 'varchar(8)', 'yyyyMMdd'],
     ['9', 'joitd', '转入日期', 'varchar(8)', ''],
     ['10', 'joimd', '转出日期', 'varchar(8)', ''],
     ['11', 'postx', '岗位', 'varchar(40)', '']])

doc.add_page_break()

# ---- PA9101 (教育) ----
h2('3.6 PA9101 — 教育经历（ZHR 自定义表）')
para('⚠️ 此表为 PEP 在 SAP 上的自定义增强表（PA9101），非 SAP 标准 PA0022。'
     'PA0022 在当前 Hive 环境中为空表（0 行），教育经历数据实际全部存储于此。')

h3('3.6.1 表级信息')
tbl(['项目', '内容'],
    [['Hive 表名', 'dmp_ods.ods_sap_pa9101'],
     ['全表行数', '2,782'],
     ['JOIN 覆盖率', '622 / 851 = 73.1%'],
     ['人均行数', '1.0（仅最高学历一条）'],
     ['总列数', '48（= 23 标准头 + 25 业务列）']])
h3('3.6.2 业务字段')
tbl(['序号', '列名', '中文名', '类型', '说明'],
    [['1', 'zhr_xuel', '学历', 'varchar(2)', '字典值：博士后/博士/硕士/本科...'],
     ['2', 'zhr_xuez', '学制', 'decimal(3,1)', '如 3.0/4.0'],
     ['3', 'zhr_xllb', '学历类别', 'varchar(2)', ''],
     ['4', 'zhr_jyzt', '教育状态', 'varchar(1)', ''],
     ['5', 'zhr_yxdm', '院校代码', 'varchar(8)', ''],
     ['6', 'zhr_yxmc', '院校名称', 'varchar(40)', ''],
     ['7', 'zhr_xkml', '学科门类', 'varchar(8)', ''],
     ['8', 'zhr_zydl', '专业大类', 'varchar(8)', ''],
     ['9', 'zhr_zymc', '专业名称', 'varchar(40)', ''],
     ['10', 'zhr_xwdj', '学位', 'varchar(1)', '字典值'],
     ['11', 'zhr_sydq', '授予地区', 'varchar(3)', ''],
     ['12', 'zhr_sydw', '授予单位', 'varchar(40)', ''],
     ['13', 'zhr_sysj', '授予时间', 'varchar(8)', 'yyyyMMdd'],
     ['14', 'zhr_xlbh', '学历证编号', 'varchar(18)', ''],
     ['15', 'zhr_xwbh', '学位证编号', 'varchar(18)', ''],
     ['16', 'zhr_sfjy', '是否就业学历', 'varchar(1)', ''],
     ['17', 'zhr_sfzg', '是否最高学历', 'varchar(1)', ''],
     ['18', 'zhr_xxxs', '学习形式', 'varchar(2)', ''],
     ['19', 'zhr_szxy', '所在学院', 'varchar(40)', ''],
     ['20', 'zhr_zmr', '证明人', 'varchar(20)', ''],
     ['21', 'zhr_cgjy', '出国境经历', 'varchar(1)', ''],
     ['22', 'zhr_dgflg', '最高学位标识', 'varchar(1)', '']])

# ---- PA9112 (资格) ----
h2('3.7 PA9112 — 资格认证（ZHR 自定义表）')
para('⚠️ 此表为 PEP 在 SAP 上的自定义增强表（PA9112），非 SAP 标准 PA0024。'
     'PA0024 在当前 Hive 环境中为空表（0 行），资格认证数据实际全部存储于此。')

h3('3.7.1 表级信息')
tbl(['项目', '内容'],
    [['Hive 表名', 'dmp_ods.ods_sap_pa9112'],
     ['全表行数', '767'],
     ['JOIN 覆盖率', '457 / 851 = 53.7%'],
     ['人均行数', '1.1（最多 5 条）'],
     ['总列数', '39（= 23 标准头 + 16 业务列）']])
h3('3.7.2 业务字段')
tbl(['序号', '列名', '中文名', '类型', '说明'],
    [['1', 'zhr_hqrq', '取得资格日期', 'varchar(8)', 'yyyyMMdd'],
     ['2', 'zhr_zgxl', '资格系列', 'varchar(8)', ''],
     ['3', 'zhr_zelb', '取得途径', 'varchar(8)', ''],
     ['4', 'zhr_zgmc', '资格名称', 'varchar(4)', '代码'],
     ['5', 'zhr_zgdj', '资格等级', 'varchar(4)', '代码'],
     ['6', 'zhr_qdtj', '取得途径明细', 'varchar(1)', ''],
     ['7', 'zhr_jdjg', '鉴定机构', 'varchar(40)', ''],
     ['8', 'zhr_gljg', '管理机构', 'varchar(40)', ''],
     ['9', 'zhr_zsbh', '证书编号', 'varchar(20)', ''],
     ['10', 'zhr_sfpg', '是否评聘', 'varchar(1)', ''],
     ['11', 'zhr_zyzg', '执业资格', 'varchar(1)', ''],
     ['12', 'zhr_zgmcwb', '资格名称（外文）', 'varchar(40)', ''],
     ['13', 'zhr_zgdjwb', '资格等级（外文）', 'varchar(20)', ''],
     ['14', 'zhr_bz', '备注', 'varchar(100)', '']])

# ---- PA0023 ----
h2('3.8 PA0023 — 工作经历')
para('记录员工的过往工作经历。一人可有多条记录。')

h3('3.8.1 表级信息')
tbl(['项目', '内容'],
    [['Hive 表名', 'dmp_ods.ods_sap_pa0023'],
     ['全表行数', '2,799'],
     ['JOIN 覆盖率', '153 / 851 = 18.0%（仅少数员工录入）'],
     ['人均行数', '1.3（最多 5 条）'],
     ['总列数', '37（= 23 标准头 + 14 业务列）']])
h3('3.8.2 业务字段')
tbl(['序号', '列名', '中文名', '类型', '说明'],
    [['1', 'arbgb', '工作单位', 'varchar(60)', ''],
     ['2', 'ort01', '城市', 'varchar(25)', ''],
     ['3', 'land1', '国家', 'varchar(3)', ''],
     ['4', 'branc', '行业', 'varchar(4)', ''],
     ['5', 'taete', '职务', 'varchar(8)', ''],
     ['6', 'deptn', '部门', 'varchar(40)', ''],
     ['7', 'jobinfo', '工作描述', 'varchar(80)', ''],
     ['8', 'refer', '证明人', 'varchar(20)', ''],
     ['9', 'refco', '证明人单位', 'varchar(40)', ''],
     ['10', 'ansvx', '工作类型', 'varchar(2)', '']])

# ---- PA0021 ----
h2('3.9 PA0021 — 家庭成员')
para('记录员工的家庭成员信息。一人可有多个家庭成员。')

h3('3.9.1 表级信息')
tbl(['项目', '内容'],
    [['Hive 表名', 'dmp_ods.ods_sap_pa0021'],
     ['全表行数', '494'],
     ['JOIN 覆盖率', '125 / 851 = 14.7%（仅少数员工录入）'],
     ['人均行数', '1.9（最多 6 人）'],
     ['总列数', '80（= 23 标准头 + 57 业务列）']])
h3('3.9.2 业务字段')
tbl(['序号', '列名', '中文名', '类型', '说明'],
    [['1', 'subty', '成员类型', 'varchar(4)', '11=父 12=母 20=配偶 30=子女'],
     ['2', 'famsa', '与员工关系', 'varchar(4)', ''],
     ['3', 'fanam', '成员姓名', 'varchar(40)', ''],
     ['4', 'fasex', '成员性别', 'varchar(1)', '1=男 2=女'],
     ['5', 'fgbdt', '成员出生日期', 'varchar(8)', 'yyyyMMdd'],
     ['6', 'fanat', '成员国籍', 'varchar(3)', ''],
     ['7', 'fasdt', '关系开始日期', 'varchar(8)', ''],
     ['8', 'erbnr', '成员证件号', 'varchar(12)', ''],
     ['9', 'ahvnr', '成员社保号', 'varchar(11)', ''],
     ['10', 'emrgn', '紧急联系人', 'varchar(1)', ''],
     ['11', 'zhr_xl', '成员学历', 'varchar(2)', ''],
     ['12', 'zhr_zzmm', '成员政治面貌', 'varchar(2)', ''],
     ['13', 'zhr_zgdw', '成员工作单位', 'varchar(40)', ''],
     ['14', 'zhr_zwzw', '成员职务', 'varchar(40)', ''],
     ['15', 'zhr_sfzh', '成员身份证号', 'varchar(20)', ''],
     ['16', 'zhr_mz', '成员民族', 'varchar(2)', ''],
     ['17', 'zhr_mebst', '成员状态', 'varchar(2)', ''],
     ['18', 'zhr_yxmc', '成员院校名称', 'varchar(40)', ''],
     ['19', 'zhr_zymc', '成员专业名称', 'varchar(40)', ''],
     ['20', 'zhr_jwork', '成员参加工作时间', 'varchar(8)', '']])

doc.add_page_break()

# ================================================================
# 4. 筛选条件与业务规则
# ================================================================
h1('4. 筛选条件与业务规则')

h2('4.1 主表筛选（PA0001）')
para('以下条件在 CTE pa1 中执行，是宽表的唯一入口筛选，确保数据口径一致：')
tbl(['序号', '条件', '字段', 'SQL', '业务含义', '类型'],
    [['1', '集团限定', 'mandt', "= '800'", 'SAP 生产集团，排除开发/测试数据', '硬规则'],
     ['2', '公司限定', 'bukrs', "= '2000'", '人民教育出版社本部，排除下属子公司', '硬规则'],
     ['3', '有效起始', 'begda', "< DATE_FORMAT(CURRENT_DATE(), 'yyyyMMdd')", '排除未来生效记录', '时效规则'],
     ['4', '有效截止', 'endda', "> DATE_FORMAT(CURRENT_DATE(), 'yyyyMMdd')", '排除已失效记录', '时效规则'],
     ['5', '在职状态', 'zhr_fjzt', "LIKE 'A%'", 'A/A10 等以 A 开头的在职/在岗状态', '业务规则']])

h2('4.2 子表筛选（通用）')
para('所有子表 CTE 均应用以下两个时效条件，确保只取"当前有效"记录：')
para('• mandt = 800 — 集团限定\n'
     '• begda < 今天 AND endda > 今天 — 有效期交叉检查')
para('额外筛选：\n'
     '• PA0185：subty = 01 — 仅取身份证\n'
     '• PA0105：PIVOT 时仅取 MAIL / CELL / MPHN 三种通讯类型')

h2('4.3 有效期处理机制')
para('SAP HR 信息类型采用"有效期"模型管理数据变更历史。每次修改产生新行，旧行 endda 被更新为修改日期：')
tbl(['场景', 'begda', 'endda', '示例'],
    [['当前在职', '20160101', '99991231', '2016年入职至今'],
     ['已离职', '20160101', '20231231', '2016年入职，2023年底离职'],
     ['岗位调动', '20240101', '99991231', '2024年调岗后新记录，旧记录 endda 变为 20231231']])
para('注：99991231 在 SAP 中表示"无限远"，即该记录当前仍然有效。')

doc.add_page_break()

# ================================================================
# 5. JOIN 逻辑与行数膨胀
# ================================================================
h1('5. JOIN 逻辑与行数膨胀')

h2('5.1 JOIN 策略总览')
tbl(['序号', '源表', 'JOIN 类型', '关联键', '粒度', '基数', '说明'],
    [['1', 'PA0001', '-', '-', '-', '851 人', '锚表，所有筛选在此执行'],
     ['2', 'PA0002', 'LEFT JOIN', 'pernr', '1:1', '848 人', '每个员工一条个人数据'],
     ['3', 'PA0105', 'LEFT JOIN', 'pernr', '1:1（透转后）', '847 人', 'PIVOT 后每人一行'],
     ['4', 'PA0185', 'LEFT JOIN', 'pernr', '1:1', '846 人', 'subty=01 每人一条'],
     ['5', 'PA0534', 'LEFT JOIN', 'pernr', '1:1', '686 人', '每人一条政治面貌'],
     ['6', 'PA9101', 'LEFT JOIN', 'pernr', '1:N', '622 人 / 1,149 行', '教育经历撑开'],
     ['7', 'PA9112', 'LEFT JOIN', 'pernr', '1:N', '457 人 / 486 行', '资格认证撑开'],
     ['8', 'PA0023', 'LEFT JOIN', 'pernr', '1:N', '153 人 / 358 行', '工作经历撑开'],
     ['9', 'PA0021', 'LEFT JOIN', 'pernr', '1:N', '125 人 / 285 行', '家庭成员撑开']])

h2('5.2 行数膨胀计算')
para('LEFT JOIN 多个 1:N 表时，输出行数为各子表行数的笛卡尔积：'
     '输出行数 ≈ 1 × N_家庭成员 × N_教育 × N_工作 × N_资格')
para('实测数据（2026-06-18）：')
tbl(['指标', '数值'],
    [['PA0001 过滤后人数', '851'],
     ['全表 JOIN 后总行数', '129,639'],
     ['膨胀倍数', '约 152 倍/人'],
     ['原因', '通讯表 PA0105 多行（每人约 157 条历史通讯记录）']])

h2('5.3 JOIN 顺序')
para('采用"锚表先行 + 子表依序 LEFT JOIN"策略。通讯表先做 PIVOT 归并为一行再 JOIN，'
     '避免通讯多行与 1:N 表产生不必要的笛卡尔积膨胀。')

doc.add_page_break()

# ================================================================
# 6. 输出字段字典
# ================================================================
h1('6. 输出字段字典')

h2('6.1 字段命名规范')
para('• 源列名：保持 SAP 原始命名（小写英文）'
     '\n• 输出别名：使用中文，按模块加前缀区分（如 教育_学历、工作_单位）'
     '\n• 1:1 字段不加前缀，1:N 字段加模块前缀以避免歧义')

h2('6.2 模块一：组织信息（来源 PA0001）')
tbl(['序号', '输出列名（中文）', '源列', '类型', '说明'],
    [['1', '人员编号', 'pernr', 'varchar(8)', '主键'],
     ['2', '姓名', 'sname', 'varchar(30)', ''],
     ['3', '英文名', 'ename', 'varchar(40)', ''],
     ['4', '公司代码', 'bukrs', 'varchar(4)', '固定 2000'],
     ['5', '人事范围', 'werks', 'varchar(4)', ''],
     ['6', '人事子范围', 'btrtl', 'varchar(4)', ''],
     ['7', '员工组', 'persg', 'varchar(1)', 'A/D/E'],
     ['8', '员工子组', 'persk', 'varchar(2)', ''],
     ['9', '组织单位', 'orgeh', 'varchar(8)', '部门'],
     ['10', '职位', 'plans', 'varchar(8)', ''],
     ['11', '职务', 'stell', 'varchar(8)', ''],
     ['12', '成本中心', 'kostl', 'varchar(10)', ''],
     ['13', '员工状态', 'zhr_ygzt', 'varchar(1)', ''],
     ['14', '附加员工状态', 'zhr_fjzt', 'varchar(3)', ''],
     ['15', '专业类型', 'zhr_zylx', 'varchar(2)', ''],
     ['16', '最高学历', 'zhr_edlst', 'varchar(2)', ''],
     ['17', '最高学位', 'zhr_dglst', 'varchar(1)', ''],
     ['18', '组织分配_开始日期', 'begda', 'varchar(8)', 'yyyyMMdd'],
     ['19', '组织分配_结束日期', 'endda', 'varchar(8)', ''],
     ['20', '组织代码', 'vdsk1', 'varchar(14)', ''],
     ['21', '工资范围', 'abkrs', 'varchar(2)', ''],
     ['22', '成本控制范围', 'kokrs', 'varchar(4)', '']])

h2('6.3 模块二：个人基础信息（来源 PA0002）')
tbl(['序号', '输出列名（中文）', '源列', '类型', '说明'],
    [['1', '姓', 'nachn', 'varchar(40)', ''],
     ['2', '名', 'vorna', 'varchar(40)', '填充率 73.1%'],
     ['3', '性别', 'gesch', 'varchar(1)', '1=男 2=女'],
     ['4', '出生日期', 'gbdat', 'varchar(8)', 'yyyyMMdd'],
     ['5', '出生地', 'gbort', 'varchar(40)', ''],
     ['6', '国籍', 'natio', 'varchar(3)', ''],
     ['7', '婚姻状况', 'famst', 'varchar(1)', '1=未婚 2=已婚'],
     ['8', '子女数', 'anzkd', 'decimal(3,0)', ''],
     ['9', '民族', 'zhr_mz', 'varchar(2)', '代码，对照数据字典'],
     ['10', '户口性质', 'zhr_hkxz', 'varchar(1)', ''],
     ['11', '籍贯省', 'zhr_jgss', 'varchar(8)', ''],
     ['12', '籍贯市', 'zhr_jgs', 'varchar(8)', ''],
     ['13', '出生省', 'zhr_csss', 'varchar(8)', ''],
     ['14', '出生市', 'zhr_css', 'varchar(8)', ''],
     ['15', '研究方向或专长', 'zhr_speci', 'varchar(40)', ''],
     ['16', '外语或少数民族语言', 'zhr_forla', 'varchar(60)', '']])

h2('6.4 模块三：通讯信息（来源 PA0105 透转）')
tbl(['序号', '输出列名（中文）', '源逻辑', '类型', '说明'],
    [['1', '邮箱', "MAX(CASE WHEN subty='MAIL' THEN usrid END)", 'varchar(30)', '841人有值'],
     ['2', '手机号', "MAX(CASE WHEN subty='CELL' THEN usrid END)", 'varchar(30)', '692人有值'],
     ['3', '座机号', "MAX(CASE WHEN subty='MPHN' THEN usrid END)", 'varchar(30)', '847人有值']])

h2('6.5 模块四～九（1:N 子表，按模块前缀分列）')
para('以下模块均为 1:N 关系，每条子记录独立一行。列名以模块前缀标识来源。')

modules = [
    ('6.5 模块四：家庭成员（来源 PA0021）', [
        ('成员类型', 'subty', '11=父 12=母 20=配偶 30=子女'),
        ('与员工关系', 'famsa', ''),
        ('成员姓名', 'fanam', ''),
        ('成员性别', 'fasex', '1=男 2=女'),
        ('成员出生日期', 'fgbdt', 'yyyyMMdd'),
        ('成员学历', 'zhr_xl', ''),
        ('成员政治面貌', 'zhr_zzmm', ''),
        ('成员工作单位', 'zhr_zgdw', ''),
        ('成员职务', 'zhr_zwzw', ''),
        ('成员身份证号', 'zhr_sfzh', ''),
        ('成员民族', 'zhr_mz', ''),
    ]),
    ('6.6 模块五：教育经历（来源 PA9101）', [
        ('教育_学历', 'zhr_xuel', '字典值'),
        ('教育_学制', 'zhr_xuez', '如 3.0/4.0'),
        ('教育_学历类别', 'zhr_xllb', ''),
        ('教育_教育状态', 'zhr_jyzt', ''),
        ('教育_院校名称', 'zhr_yxmc', ''),
        ('教育_学科门类', 'zhr_xkml', ''),
        ('教育_专业名称', 'zhr_zymc', ''),
        ('教育_学位', 'zhr_xwdj', '字典值'),
        ('教育_授予单位', 'zhr_sydw', ''),
        ('教育_授予时间', 'zhr_sysj', 'yyyyMMdd'),
        ('教育_学历证编号', 'zhr_xlbh', ''),
        ('教育_学位证编号', 'zhr_xwbh', ''),
        ('教育_是否就业学历', 'zhr_sfjy', ''),
        ('教育_是否最高学历', 'zhr_sfzg', ''),
        ('教育_学习形式', 'zhr_xxxs', ''),
        ('教育_所在学院', 'zhr_szxy', ''),
        ('教育_证明人', 'zhr_zmr', ''),
        ('教育_是否出国境经历', 'zhr_cgjy', ''),
    ]),
    ('6.7 模块六：工作经历（来源 PA0023）', [
        ('工作_工作单位', 'arbgb', ''),
        ('工作_城市', 'ort01', ''),
        ('工作_职务', 'taete', ''),
        ('工作_部门', 'deptn', ''),
        ('工作_工作描述', 'jobinfo', ''),
        ('工作_证明人', 'refer', ''),
        ('工作_开始日期', 'begda', 'yyyyMMdd'),
        ('工作_结束日期', 'endda', 'yyyyMMdd'),
    ]),
    ('6.8 模块七：资格认证（来源 PA9112）', [
        ('资格_取得日期', 'zhr_hqrq', 'yyyyMMdd'),
        ('资格_资格名称', 'zhr_zgmc', ''),
        ('资格_资格等级', 'zhr_zgdj', ''),
        ('资格_鉴定机构', 'zhr_jdjg', ''),
        ('资格_管理机构', 'zhr_gljg', ''),
        ('资格_证书编号', 'zhr_zsbh', ''),
    ]),
    ('6.9 模块八：证件信息（来源 PA0185）', [
        ('证件号码', 'icnum', '身份证号'),
        ('签发机构', 'auth1', ''),
        ('证件有效期开始', 'usefr', 'yyyyMMdd'),
        ('证件有效期截至', 'useto', 'yyyyMMdd'),
    ]),
    ('6.10 模块九：政治面貌（来源 PA0534）', [
        ('政治面貌代码', 'pcode', ''),
        ('加入组织名称', 'joinu', ''),
        ('加入日期', 'joipd', 'yyyyMMdd'),
        ('转正日期', 'bfmdt', 'yyyyMMdd'),
    ]),
]
for title, fields in modules:
    h3(title)
    tbl(['输出列名（中文）', '源列', '说明'], fields)

doc.add_page_break()

# ================================================================
# 7. 数据质量评估
# ================================================================
h1('7. 数据质量评估')

h2('7.1 JOIN 覆盖率评级')
tbl(['源表', '匹配人数/总人数', '覆盖率', '评级', '说明'],
    [['PA0002', '848 / 851', '99.6%', '🟢 优秀', '仅 3 人缺失个人数据'],
     ['PA0105', '847 / 851', '99.5%', '🟢 优秀', '仅 4 人缺失通讯数据'],
     ['PA0185', '846 / 851', '99.4%', '🟢 优秀', '仅 5 人缺失身份证'],
     ['PA0534', '686 / 851', '80.6%', '🟡 良好', '165 人无政治面貌，非党员群体正常'],
     ['PA9101', '622 / 851', '73.1%', '🟡 良好', '229 人无教育记录，可能为低学历/未录入'],
     ['PA9112', '457 / 851', '53.7%', '🟠 注意', '394 人无资格认证，可能不需要或未录入'],
     ['PA0023', '153 / 851', '18.0%', '🔴 低', '698 人无工作经历，字段非必填'],
     ['PA0021', '125 / 851', '14.7%', '🔴 低', '726 人无家庭成员，字段非必填']])

h2('7.2 字段级缺失清单')
tbl(['表', '字段', '中文名', '当前填充率', '影响评估', '建议'],
    [['PA0001', 'ansvh', '工作合同', '0%', '低 — 该字段 SAP 未启用', '可忽略'],
     ['PA0001', 'zhr_edlst', '最高学历', '61.5%', '中 — 38.5%员工缺学历', '关联 PA9101 补全'],
     ['PA0001', 'zhr_dglst', '最高学位', '61.3%', '中 — 38.7%员工缺学位', '关联 PA9101 补全'],
     ['PA0002', 'vorna', '名', '73.1%', '低 — 179人缺名', '确认是否外籍员工'],
     ['PA0002', 'famst', '婚姻状况', '19.2%', '中 — 80%为空', '确认 SAP 是否维护'],
     ['PA0002', 'zhr_sthea', '健康状况', '0%', '中 — 全空', '确认 SAP 是否维护'],
     ['PA0002', 'zhr_relig', '宗教信仰', '0%', '低 — 非必要', '可忽略'],
     ['PA0002', 'zhr_heduc', '最高学历', '0.8%', '低 — 主数据在 PA0001', '使用 PA0001 字段']])

h2('7.3 维度分布快照')
tbl(['维度', '分布'],
    [['性别', '男 293 (34.4%) | 女 555 (65.2%) | 未知 3'],
     ['员工组', 'A10=292 | A20=455 | A40=30 | E20=27 | D20=21 | E10=16 | 其他=10'],
     ['婚姻', '未婚 86 (10.1%) | 空值 685 (80.5%) | 其他 80'],
     ['民族 TOP5', '00=汉族 786 | 10=23 | 02=9 | 01=5 | 09=5'],
     ['公司代码', '2000=851（筛选条件固定）']])

doc.add_page_break()

# ================================================================
# 8. 完整 SQL
# ================================================================
h1('8. 完整 SQL')

para('以下为员工主数据宽表的完整查询。可直接在 Hive / Spark SQL 环境执行。')
para('• 目标库：dmp_ods（只读）\n'
     '• SQL 方言：Spark SQL（Hive 兼容）\n'
     '• 执行方式：dbt run / spark-sql / beeline')

code_block("""-- ============================================================================
-- PEP HR 员工主数据宽表
-- 数据源：SAP HR 信息类型 → Hive dmp_ods
-- 范围：集团 800 | 公司 2000 | 在职员工（zhr_fjzt LIKE 'A%'）
-- 日期基准：CURRENT_DATE()
-- ============================================================================

WITH
-- ====== 1. 组织分配（锚表） ======
pa1 AS (
    SELECT *
    FROM dmp_ods.ods_sap_pa0001
    WHERE mandt = '800'
      AND bukrs = '2000'
      AND begda < date_format(current_date(), 'yyyyMMdd')
      AND endda > date_format(current_date(), 'yyyyMMdd')
      AND zhr_fjzt LIKE 'A%'
),

-- ====== 2. 个人数据 ======
pa2 AS (
    SELECT *
    FROM dmp_ods.ods_sap_pa0002
    WHERE mandt = '800'
      AND begda < date_format(current_date(), 'yyyyMMdd')
      AND endda > date_format(current_date(), 'yyyyMMdd')
),

-- ====== 3. 通讯信息（行转列） ======
pa0105_pivot AS (
    SELECT
        pernr,
        MAX(CASE WHEN subty = 'MAIL' THEN usrid END) AS 邮箱,
        MAX(CASE WHEN subty = 'CELL' THEN usrid END) AS 手机号,
        MAX(CASE WHEN subty = 'MPHN' THEN usrid END) AS 座机号
    FROM dmp_ods.ods_sap_pa0105
    WHERE mandt = '800'
      AND begda < date_format(current_date(), 'yyyyMMdd')
      AND endda > date_format(current_date(), 'yyyyMMdd')
    GROUP BY pernr
),

-- ====== 4. 证件信息（仅身份证） ======
pa0185 AS (
    SELECT *
    FROM dmp_ods.ods_sap_pa0185
    WHERE mandt = '800'
      AND begda < date_format(current_date(), 'yyyyMMdd')
      AND endda > date_format(current_date(), 'yyyyMMdd')
      AND subty = '01'
),

-- ====== 5. 政治面貌 ======
pa0534 AS (
    SELECT *
    FROM dmp_ods.ods_sap_pa0534
    WHERE mandt = '800'
      AND begda < date_format(current_date(), 'yyyyMMdd')
      AND endda > date_format(current_date(), 'yyyyMMdd')
),

-- ====== 6. 教育经历（ZHR PA9101） ======
pa9101 AS (
    SELECT *
    FROM dmp_ods.ods_sap_pa9101
    WHERE mandt = '800'
      AND begda < date_format(current_date(), 'yyyyMMdd')
      AND endda > date_format(current_date(), 'yyyyMMdd')
),

-- ====== 7. 资格认证（ZHR PA9112） ======
pa9112 AS (
    SELECT *
    FROM dmp_ods.ods_sap_pa9112
    WHERE mandt = '800'
      AND begda < date_format(current_date(), 'yyyyMMdd')
      AND endda > date_format(current_date(), 'yyyyMMdd')
),

-- ====== 8. 工作经历 ======
pa23 AS (
    SELECT *
    FROM dmp_ods.ods_sap_pa0023
    WHERE mandt = '800'
      AND begda < date_format(current_date(), 'yyyyMMdd')
      AND endda > date_format(current_date(), 'yyyyMMdd')
),

-- ====== 9. 家庭成员 ======
pa21 AS (
    SELECT *
    FROM dmp_ods.ods_sap_pa0021
    WHERE mandt = '800'
      AND begda < date_format(current_date(), 'yyyyMMdd')
      AND endda > date_format(current_date(), 'yyyyMMdd')
)

-- ====== 主查询 ======
SELECT
    -- 模块一：组织信息（PA0001）--
    p1.pernr                                AS 人员编号,
    p1.sname                                AS 姓名,
    p1.ename                                AS 英文名,
    p1.bukrs                                AS 公司代码,
    p1.werks                                AS 人事范围,
    p1.btrtl                                AS 人事子范围,
    p1.persg                                AS 员工组,
    p1.persk                                AS 员工子组,
    p1.orgeh                                AS 组织单位,
    p1.plans                                AS 职位,
    p1.stell                                AS 职务,
    p1.kostl                                AS 成本中心,
    p1.zhr_ygzt                             AS 员工状态,
    p1.zhr_fjzt                             AS 附加员工状态,
    p1.zhr_zylx                             AS 专业类型,
    p1.zhr_edlst                            AS 最高学历,
    p1.zhr_dglst                            AS 最高学位,
    p1.begda                                AS 组织分配_开始日期,
    p1.endda                                AS 组织分配_结束日期,
    p1.vdsk1                                AS 组织代码,
    p1.abkrs                                AS 工资范围,
    p1.kokrs                                AS 成本控制范围,

    -- 模块二：个人基础信息（PA0002）--
    p2.nachn                                AS 姓,
    p2.vorna                                AS 名,
    p2.gesch                                AS 性别,
    p2.gbdat                                AS 出生日期,
    p2.gbort                                AS 出生地,
    p2.natio                                AS 国籍,
    p2.famst                                AS 婚姻状况,
    p2.anzkd                                AS 子女数,
    p2.zhr_mz                               AS 民族,
    p2.zhr_hkxz                             AS 户口性质,
    p2.zhr_jgss                             AS 籍贯省,
    p2.zhr_jgs                              AS 籍贯市,
    p2.zhr_csss                             AS 出生省,
    p2.zhr_css                              AS 出生市,
    p2.zhr_speci                            AS 研究方向或专长,
    p2.zhr_forla                            AS 外语或少数民族语言,

    -- 模块三：通讯信息（PA0105 透转）--
    c.邮箱,
    c.手机号,
    c.座机号,

    -- 模块四：家庭成员（PA0021）--
    f.subty                                 AS 成员类型,
    f.famsa                                 AS 与员工关系,
    f.fanam                                 AS 成员姓名,
    f.fasex                                 AS 成员性别,
    f.fgbdt                                 AS 成员出生日期,
    f.zhr_xl                                AS 成员学历,
    f.zhr_zzmm                              AS 成员政治面貌,
    f.zhr_zgdw                              AS 成员工作单位,
    f.zhr_zwzw                              AS 成员职务,
    f.zhr_sfzh                              AS 成员身份证号,
    f.zhr_mz                                AS 成员民族,

    -- 模块五：教育经历（PA9101）--
    e9101.zhr_xuel                          AS 教育_学历,
    e9101.zhr_xuez                          AS 教育_学制,
    e9101.zhr_xllb                          AS 教育_学历类别,
    e9101.zhr_jyzt                          AS 教育_教育状态,
    e9101.zhr_yxmc                          AS 教育_院校名称,
    e9101.zhr_xkml                          AS 教育_学科门类,
    e9101.zhr_zydl                          AS 教育_专业大类,
    e9101.zhr_zymc                          AS 教育_专业名称,
    e9101.zhr_xwdj                          AS 教育_学位,
    e9101.zhr_sydq                          AS 教育_授予地区,
    e9101.zhr_sydw                          AS 教育_授予单位,
    e9101.zhr_sysj                          AS 教育_授予时间,
    e9101.zhr_xlbh                          AS 教育_学历证编号,
    e9101.zhr_xwbh                          AS 教育_学位证编号,
    e9101.zhr_sfjy                          AS 教育_是否就业学历,
    e9101.zhr_sfzg                          AS 教育_是否最高学历,
    e9101.zhr_xxxs                          AS 教育_学习形式,
    e9101.zhr_szxy                          AS 教育_所在学院,
    e9101.zhr_zmr                           AS 教育_证明人,
    e9101.zhr_cgjy                          AS 教育_是否出国境经历,

    -- 模块六：工作经历（PA0023）--
    w.arbgb                                 AS 工作_工作单位,
    w.ort01                                 AS 工作_城市,
    w.land1                                 AS 工作_国家,
    w.taete                                 AS 工作_职务,
    w.deptn                                 AS 工作_部门,
    w.jobinfo                               AS 工作_工作描述,
    w.refer                                 AS 工作_证明人,
    w.begda                                 AS 工作_开始日期,
    w.endda                                 AS 工作_结束日期,

    -- 模块七：资格认证（PA9112）--
    q9112.zhr_hqrq                          AS 资格_取得日期,
    q9112.zhr_zgxl                          AS 资格_资格系列,
    q9112.zhr_zelb                          AS 资格_取得途径,
    q9112.zhr_zgmc                          AS 资格_资格名称,
    q9112.zhr_zgdj                          AS 资格_资格等级,
    q9112.zhr_qdtj                          AS 资格_取得途径明细,
    q9112.zhr_jdjg                          AS 资格_鉴定机构,
    q9112.zhr_gljg                          AS 资格_管理机构,
    q9112.zhr_zsbh                          AS 资格_证书编号,
    q9112.zhr_sfpg                          AS 资格_是否评聘,
    q9112.zhr_zyzg                          AS 资格_执业资格,
    q9112.zhr_bz                            AS 资格_备注,

    -- 模块八：证件信息（PA0185）--
    i.icnum                                 AS 证件号码,
    i.auth1                                 AS 签发机构,
    i.isspl                                 AS 发证机关,
    i.usefr                                 AS 证件有效期开始,
    i.useto                                 AS 证件有效期截至,

    -- 模块九：政治面貌（PA0534）--
    pol.pcode                               AS 政治面貌代码,
    pol.joinu                               AS 加入组织名称,
    pol.intr1                               AS 介绍人1,
    pol.intr2                               AS 介绍人2,
    pol.paety                               AS 党员类型,
    pol.pstat                               AS 状态,
    pol.joipd                               AS 加入日期,
    pol.bfmdt                               AS 转正日期,
    pol.joitd                               AS 转入日期,
    pol.joimd                               AS 转出日期,
    pol.postx                               AS 岗位

FROM pa1 p1
LEFT JOIN pa2 p2         ON p1.pernr = p2.pernr
LEFT JOIN pa0105_pivot c ON p1.pernr = c.pernr
LEFT JOIN pa0185 i       ON p1.pernr = i.pernr
LEFT JOIN pa0534 pol     ON p1.pernr = pol.pernr
LEFT JOIN pa9101 e9101   ON p1.pernr = e9101.pernr
LEFT JOIN pa9112 q9112   ON p1.pernr = q9112.pernr
LEFT JOIN pa23 w         ON p1.pernr = w.pernr
LEFT JOIN pa21 f         ON p1.pernr = f.pernr
""")

doc.add_page_break()

# ================================================================
# 9. 附录
# ================================================================
h1('9. 附录')

h2('9.1 注意事项')
para('1. PA0105 列名全大写且 varchar(512)：在宽表 SQL 中引用时注意大小写\n'
     '2. PA9101/PA9112 为 ZHR 自定义表：SAP 标准 PA0022/PA0024 为空表，切换使用\n'
     '3. 日期字段为 varchar 存储：yyyyMMdd 格式，字典序比较与日期序一致，无需 CAST\n'
     '4. 99991231 表示永久有效：在 endda 比较中 > 当前日期即有效\n'
     '5. 行数膨胀：含 1:N 表时结果集远超人数，BI 消费时需注意聚合粒度')

h2('9.2 数据字典映射')
para('民族代码（PA0002.zhr_mz / PA0021.zhr_mz）→ 参见《人资数据库表及数据字典整理.xlsx》Sheet2 民族列\n'
     '学历代码（PA0001.zhr_edlst / PA9101.zhr_xuel）→ 参见同文件 Sheet2 学历列\n'
     '学位代码（PA0001.zhr_dglst / PA9101.zhr_xwdj）→ 参见同文件 Sheet2 学位列\n'
     '注：Sheet2 中学位列标注"已和人资确认不需要"，但 PA0001 中仍存储了学位代码')

h2('9.3 未纳入的表')
tbl(['表名', 'SAP 标准用途', '当前状态', '未纳入原因'],
    [['ods_sap_pa0022', '教育经历', '空表（0行）', '数据在 PA9101'],
     ['ods_sap_pa0024', '资格认证', '空表（0行）', '数据在 PA9112'],
     ['ods_sap_pa105', '通讯（冗余）', '15,897行', '与 PA0105 重复，待确认后废弃']])

# ── 保存 ──
out = os.path.expanduser(r'~\Desktop\PEP_HR员工主数据宽表_技术规格说明书.docx')
doc.save(out)
print(f'Done: {out}')
