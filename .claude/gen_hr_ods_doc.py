# -*- coding: utf-8 -*-
"""生成 HR ODS 层表结构 Word 文档"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── 页面设置 ──
section = doc.sections[0]
section.page_width = Cm(29.7)  # A4 横向
section.page_height = Cm(21.0)

# ── 样式 ──
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(9)
style.paragraph_format.space_after = Pt(2)

# ============================================================
# 标题
# ============================================================
title = doc.add_heading('PEP 数据仓库 — HR 人资模块 ODS 层表结构', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('数据库：dmp_ods | 数据源：SAP HR 信息类型 | 日期：2026-06-18')
doc.add_paragraph('')

# ============================================================
# 一、总览
# ============================================================
doc.add_heading('一、ODS 层 HR 表总览', level=1)

overview_data = [
    ['表名', 'SAP 信息类型', '行数', '列数', '状态', '对应 Excel 模块'],
    ['ods_sap_pa0001', 'PA0001 组织分配', '12,318', '57', '有数据', '基本信息(组织)'],
    ['ods_sap_pa0002', 'PA0002 个人数据', '5,286', '87', '有数据', '基本信息(个人)'],
    ['ods_sap_pa0021', 'PA0021 家庭成员', '494', '80', '有数据', '家庭成员'],
    ['ods_sap_pa0022', 'PA0022 教育经历', '0', '43', '⚠ 空表', '教育经历'],
    ['ods_sap_pa0023', 'PA0023 工作经历', '2,799', '37', '有数据', '工作经历'],
    ['ods_sap_pa0024', 'PA0024 资格认证', '0', '25', '⚠ 空表', '资格认证'],
    ['ods_sap_pa0105', 'PA0105 通讯', '381,133', '26', '有数据', '基本信息(通讯)'],
    ['ods_sap_pa0185', 'PA0185 证件', '5,212', '43', '有数据', '证件信息'],
    ['ods_sap_pa0534', 'PA0534 政治面貌', '3,351', '40', '有数据', '政治面貌'],
    ['ods_sap_pa105', '(PA0105 重复)', '15,897', '26', '⚠ 冗余表', '—'],
]

table = doc.add_table(rows=len(overview_data), cols=6, style='Table Grid')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, row_data in enumerate(overview_data):
    row = table.rows[i]
    for j, cell_text in enumerate(row_data):
        cell = row.cells[j]
        cell.text = cell_text
        if i == 0:
            for p in cell.paragraphs:
                p.runs[0].bold = True

doc.add_paragraph('')

# ============================================================
# 二、SAP HR 信息类型标准头（所有表共有）
# ============================================================
doc.add_heading('二、SAP HR 信息类型标准头（22列，所有表共有）', level=1)

doc.add_paragraph(
    '以下 22 列为 SAP HR 信息类型的标准管理字段，每张 PA 表均包含。'
    '（注：ods_sap_pa0105 例外，列名全大写且数据类型为 varchar(512)）'
)

header_fields = [
    ['列名', '数据类型', '说明', '备注'],
    ['mandt', 'varchar(3)', '集团', '固定 800'],
    ['pernr', 'varchar(8)', '人员编号', '主键'],
    ['subty', 'varchar(4)', '子信息类型', '如 MAIL/CELL/MPHN/01'],
    ['objps', 'varchar(2)', '对象标识', ''],
    ['sprps', 'varchar(1)', '锁定标识符', ''],
    ['endda', 'varchar(8)', '结束日期', '99991231 = 当前有效'],
    ['begda', 'varchar(8)', '开始日期', 'yyyyMMdd 格式'],
    ['seqnr', 'varchar(3)', '信息类型记录号', '复合主键之一'],
    ['aedtm', 'varchar(8)', '更改日期', ''],
    ['uname', 'varchar(12)', '更改者', '如 HR01'],
    ['histo', 'varchar(1)', '历史记录', ''],
    ['itxex', 'varchar(1)', '文本存在', ''],
    ['refex', 'varchar(1)', '参考存在', ''],
    ['ordex', 'varchar(1)', '确定字段存在', ''],
    ['itbld', 'varchar(2)', '屏幕控制', ''],
    ['preas', 'varchar(2)', '更改原因', ''],
    ['flag1', 'varchar(1)', '保留字段1', ''],
    ['flag2', 'varchar(1)', '保留字段2', ''],
    ['flag3', 'varchar(1)', '保留字段3', ''],
    ['flag4', 'varchar(1)', '保留字段4', ''],
    ['rese1', 'varchar(2)', '保留字段5', ''],
    ['rese2', 'varchar(2)', '保留字段6', ''],
    ['grpvl', 'varchar(4)', '分组值', ''],
]

table = doc.add_table(rows=len(header_fields), cols=4, style='Table Grid')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, row_data in enumerate(header_fields):
    row = table.rows[i]
    for j, cell_text in enumerate(row_data):
        cell = row.cells[j]
        cell.text = cell_text
        if i == 0:
            for p in cell.paragraphs:
                p.runs[0].bold = True

doc.add_paragraph('')

# ============================================================
# 三、各表业务字段
# ============================================================
doc.add_heading('三、各表业务字段详情', level=1)

tables_detail = [
    {
        'name': 'ods_sap_pa0001',
        'desc': '人力资源主记录：信息类型 0001（组织分配）',
        'rows': '12,318',
        'cols': '57',
        'fields': [
            ['bukrs', 'varchar(4)', '公司代码', '如 1000/2021'],
            ['werks', 'varchar(4)', '人事范围', '如 1000/2021'],
            ['persg', 'varchar(1)', '员工组', 'A/B/C...'],
            ['persk', 'varchar(2)', '员工子组', '10/40...'],
            ['vdsk1', 'varchar(14)', '组织代码', ''],
            ['gsber', 'varchar(4)', '业务范围', ''],
            ['btrtl', 'varchar(4)', '人事子范围', '如 1101'],
            ['juper', 'varchar(4)', '法律', ''],
            ['abkrs', 'varchar(2)', '工资范围', ''],
            ['ansvh', 'varchar(2)', '工作合同', ''],
            ['kostl', 'varchar(10)', '成本中心', '如 1000060000'],
            ['orgeh', 'varchar(8)', '组织单位', '如 10000039'],
            ['plans', 'varchar(8)', '职位', '如 20000034'],
            ['stell', 'varchar(8)', '职务', ''],
            ['mstbr', 'varchar(8)', '主管', ''],
            ['sacha', 'varchar(3)', '工资核算管理员', ''],
            ['sachp', 'varchar(3)', '人事管理员', ''],
            ['sachz', 'varchar(3)', '时间管理员', ''],
            ['sname', 'varchar(30)', '姓名', ''],
            ['ename', 'varchar(40)', '英文名', ''],
            ['otype', 'varchar(2)', '对象类型', 'S=员工'],
            ['sbmod', 'varchar(4)', '组', ''],
            ['kokrs', 'varchar(4)', '成本控制范围', '如 CEPM'],
            ['fistl', 'varchar(16)', '资金中心', ''],
            ['geber', 'varchar(10)', '基金', ''],
            ['fkber', 'varchar(16)', '功能范围', ''],
            ['grant_nbr', 'varchar(20)', '同意', ''],
            ['sgmnt', 'varchar(10)', '段', ''],
            ['budget_pd', 'varchar(10)', '预算期间', ''],
            ['zhr_ygzt', 'varchar(1)', '员工状态', 'A=在职'],
            ['zhr_fjzt', 'varchar(3)', '附加员工状态', '如 A10'],
            ['zhr_zylx', 'varchar(2)', '职业类型', '如 03'],
            ['zhr_edlst', 'varchar(2)', '最高学历', ''],
            ['zhr_dglst', 'varchar(1)', '最高学位', ''],
        ]
    },
    {
        'name': 'ods_sap_pa0002',
        'desc': '人力资源主记录：信息类型 0002（个人数据）',
        'rows': '5,286',
        'cols': '87',
        'fields': [
            ['inits', 'varchar(10)', '首字母', ''],
            ['nachn', 'varchar(40)', '姓', ''],
            ['name2', 'varchar(40)', '名2', ''],
            ['nach2', 'varchar(40)', '姓2', ''],
            ['vorna', 'varchar(40)', '名', ''],
            ['cname', 'varchar(80)', '中文姓名', ''],
            ['titel', 'varchar(15)', '头衔', ''],
            ['titl2', 'varchar(15)', '头衔2', ''],
            ['namzu', 'varchar(15)', '名称前缀', ''],
            ['vorsw', 'varchar(15)', '名前缀', ''],
            ['vors2', 'varchar(15)', '名前缀2', ''],
            ['rufnm', 'varchar(40)', '称呼名', ''],
            ['midnm', 'varchar(40)', '中间名', ''],
            ['knznm', 'varchar(2)', '名标识', ''],
            ['anred', 'varchar(1)', '称谓', ''],
            ['gesch', 'varchar(1)', '性别', '1=男 2=女'],
            ['gbdat', 'varchar(8)', '出生日期', 'yyyyMMdd'],
            ['gblnd', 'varchar(3)', '出生国', ''],
            ['gbdep', 'varchar(3)', '出生省', ''],
            ['gbort', 'varchar(40)', '出生地', ''],
            ['natio', 'varchar(3)', '国籍', ''],
            ['nati2', 'varchar(3)', '国籍2', ''],
            ['nati3', 'varchar(3)', '国籍3', ''],
            ['sprsl', 'varchar(1)', '语言', ''],
            ['konfe', 'varchar(2)', '宗教', ''],
            ['famst', 'varchar(1)', '婚姻状况', '1=未婚 2=已婚'],
            ['famdt', 'varchar(8)', '结婚日期', ''],
            ['anzkd', 'decimal(3,0)', '子女数', ''],
            ['nacon', 'varchar(1)', '国籍标识', ''],
            ['permo', 'varchar(2)', '工资周期', ''],
            ['perid', 'varchar(20)', '身份证号', ''],
            ['gbpas', 'varchar(8)', '出生日期(旧)', ''],
            ['fnamk', 'varchar(40)', '父姓', ''],
            ['lnamk', 'varchar(40)', '父名', ''],
            ['fnamr', 'varchar(40)', '母姓', ''],
            ['lnamr', 'varchar(40)', '母名', ''],
            ['nabik', 'varchar(40)', '配偶姓', ''],
            ['nabir', 'varchar(40)', '配偶名', ''],
            ['nickk', 'varchar(40)', '昵称', ''],
            ['nickr', 'varchar(40)', '昵称2', ''],
            ['gbjhr', 'varchar(4)', '出生年', ''],
            ['gbmon', 'varchar(2)', '出生月', ''],
            ['gbtag', 'varchar(2)', '出生日', ''],
            ['nchmc', 'varchar(25)', '中文姓', ''],
            ['vnamc', 'varchar(25)', '中文名', ''],
            ['namz2', 'varchar(15)', '名称后缀2', ''],
            # ZHR 自定义字段
            ['zhr_mz', 'varchar(2)', '民族', ''],
            ['zhr_hkxz', 'varchar(1)', '户口性质', ''],
            ['zhr_jgss', 'varchar(8)', '籍贯省', ''],
            ['zhr_jgs', 'varchar(8)', '籍贯市', ''],
            ['zhr_csss', 'varchar(8)', '出生省', ''],
            ['zhr_css', 'varchar(8)', '出生市', ''],
            ['zhr_tssf', 'varchar(1)', '特殊身份', ''],
            ['zhr_dasf', 'varchar(1)', '档案身份', ''],
            ['hyguid', 'varchar(32)', '唯一标识', ''],
            ['zhr_jgq', 'varchar(8)', '籍贯区', ''],
            ['zhr_csq', 'varchar(8)', '出生区', ''],
            ['zhr_sthea', 'varchar(40)', '健康状况', ''],
            ['zhr_relig', 'varchar(40)', '宗教信仰', ''],
            ['zhr_speci', 'varchar(40)', '研究方向/专长', ''],
            ['zhr_forla', 'varchar(60)', '外语/少数民族语言', ''],
            ['zhr_wybrk', 'varchar(255)', '外语说明', ''],
            ['zhr_heduc', 'varchar(2)', '最高学历', ''],
            ['zhr_hdegr', 'varchar(1)', '最高学位', ''],
        ]
    },
    {
        'name': 'ods_sap_pa0021',
        'desc': '人力资源主记录：信息类型 0021（家庭成员）',
        'rows': '494',
        'cols': '80',
        'fields': [
            ['famsa', 'varchar(4)', '成员关系', '11=父 12=母 20=配偶 30=子女'],
            ['fgbdt', 'varchar(8)', '成员出生日期', ''],
            ['fgbld', 'varchar(3)', '出生国', ''],
            ['fanat', 'varchar(3)', '国籍', ''],
            ['fasex', 'varchar(1)', '成员性别', '1=男 2=女'],
            ['favor', 'varchar(40)', '成员名', ''],
            ['fanam', 'varchar(40)', '成员姓名', ''],
            ['fgbot', 'varchar(40)', '出生地', ''],
            ['fgdep', 'varchar(3)', '出生省', ''],
            ['erbnr', 'varchar(12)', '证件号', ''],
            ['fgbna', 'varchar(40)', '成员姓', ''],
            ['fnac2', 'varchar(40)', '成员名2', ''],
            ['fcnam', 'varchar(80)', '中文名', ''],
            ['fknzn', 'varchar(2)', '名标识', ''],
            ['finit', 'varchar(10)', '首字母', ''],
            ['fvrsw', 'varchar(15)', '名前缀', ''],
            ['fvrs2', 'varchar(15)', '名前缀2', ''],
            ['fnmzu', 'varchar(15)', '名称前缀', ''],
            ['ahvnr', 'varchar(11)', '社保号', ''],
            ['kdsvh', 'varchar(2)', '健康保险', ''],
            ['kdbsl', 'varchar(2)', '保险类型', ''],
            ['kdutb', 'varchar(2)', '保险机构', ''],
            ['kdgbr', 'varchar(2)', '出生日期(保险)', ''],
            ['kdart', 'varchar(2)', '保险种类', ''],
            ['kdzug', 'varchar(2)', '入保日期', ''],
            ['kdzul', 'varchar(2)', '离保日期', ''],
            ['kdvbe', 'varchar(2)', '保险备注', ''],
            ['ermnr', 'varchar(8)', '成员编号', ''],
            ['ausvl', 'varchar(4)', '居住状态', ''],
            ['ausvg', 'varchar(8)', '居住期限', ''],
            ['fasdt', 'varchar(8)', '关系开始日期', ''],
            ['fasar', 'varchar(2)', '成员状态', ''],
            ['fasin', 'varchar(20)', '身份证号', ''],
            ['egaga', 'varchar(8)', '配偶收入', ''],
            ['fana2', 'varchar(3)', '国籍2', ''],
            ['fana3', 'varchar(3)', '国籍3', ''],
            ['betrg', 'decimal(9,2)', '金额', ''],
            ['titel', 'varchar(15)', '头衔', ''],
            ['emrgn', 'varchar(1)', '紧急联系人', ''],
            ['zhr_xl', 'varchar(2)', '学历', ''],
            ['zhr_zzmm', 'varchar(2)', '政治面貌', ''],
            ['zhr_zgdw', 'varchar(40)', '工作单位', ''],
            ['zhr_zwzw', 'varchar(40)', '职务', ''],
            ['zhr_dz', 'varchar(60)', '地址', ''],
            ['zhr_lxdh', 'varchar(20)', '联系电话', ''],
            ['zhr_sfzh', 'varchar(20)', '身份证号', ''],
            ['hrguid', 'varchar(32)', '唯一标识', ''],
            ['zhr_wrkst', 'varchar(2)', '工作情况', ''],
            ['zhr_mebst', 'varchar(2)', '成员状态', ''],
            ['zhr_mz', 'varchar(2)', '民族', ''],
            ['zhr_zyjszw', 'varchar(40)', '专业技术职务', ''],
            ['zhr_yxmc', 'varchar(40)', '院校名称', ''],
            ['zhr_zymc', 'varchar(40)', '专业名称', ''],
            ['zhr_jwork', 'varchar(8)', '参加工作时间', ''],
            ['zhr_jgss', 'varchar(8)', '籍贯省', ''],
            ['zhr_jgs', 'varchar(8)', '籍贯市', ''],
            ['zhr_jgq', 'varchar(8)', '籍贯区', ''],
        ]
    },
    {
        'name': 'ods_sap_pa0022',
        'desc': '人力资源主记录：信息类型 0022（教育经历）⚠ 空表，0 行',
        'rows': '0',
        'cols': '43',
        'fields': [
            ['slart', 'varchar(2)', '教育类型', ''],
            ['insti', 'varchar(80)', '毕业院校', ''],
            ['sland', 'varchar(3)', '国家', ''],
            ['ausbi', 'varchar(8)', '学历', ''],
            ['slabs', 'varchar(2)', '学制', ''],
            ['anzkl', 'decimal(3,0)', '年限', ''],
            ['anzeh', 'varchar(3)', '学时单位', ''],
            ['sltp1', 'varchar(5)', '学历类别', ''],
            ['sltp2', 'varchar(5)', '学历类别2', ''],
            ['jbez1', 'decimal(11,2)', '学费', ''],
            ['waers', 'varchar(5)', '货币', ''],
            ['slpln', 'varchar(1)', '计划', ''],
            ['slktr', 'varchar(2)', '教育状态', ''],
            ['slrzg', 'varchar(1)', '学位', ''],
            ['ksbez', 'varchar(25)', '证书编号', ''],
            ['tx122', 'varchar(40)', '证明人', ''],
            ['schcd', 'varchar(10)', '学校代码', ''],
            ['faccd', 'varchar(3)', '院系', ''],
            ['dptmt', 'varchar(40)', '专业', ''],
            ['emark', 'varchar(4)', '是否就业学历', ''],
        ]
    },
    {
        'name': 'ods_sap_pa0023',
        'desc': '人力资源主记录：信息类型 0023（工作经历）',
        'rows': '2,799',
        'cols': '37',
        'fields': [
            ['arbgb', 'varchar(60)', '工作单位', ''],
            ['ort01', 'varchar(25)', '城市', ''],
            ['land1', 'varchar(3)', '国家', ''],
            ['branc', 'varchar(4)', '行业', ''],
            ['taete', 'varchar(8)', '职务', ''],
            ['ansvx', 'varchar(2)', '工作类型', ''],
            ['ortj1', 'varchar(40)', '地址1', ''],
            ['ortj2', 'varchar(40)', '地址2', ''],
            ['ortj3', 'varchar(40)', '地址3', ''],
            ['deptn', 'varchar(40)', '部门', ''],
            ['jobinfo', 'varchar(80)', '工作描述', ''],
            ['refer', 'varchar(20)', '证明人', ''],
            ['refco', 'varchar(40)', '证明人单位', ''],
            ['hrguid', 'varchar(32)', '唯一标识', ''],
        ]
    },
    {
        'name': 'ods_sap_pa0024',
        'desc': '人力资源主记录：信息类型 0024（资格认证）⚠ 空表，0 行',
        'rows': '0',
        'cols': '25',
        'fields': [
            ['quali', 'varchar(8)', '资格名称', ''],
            ['auspr', 'varchar(4)', '资格等级', ''],
        ]
    },
    {
        'name': 'ods_sap_pa0105',
        'desc': '人力资源主记录：信息类型 0105（通讯）⚠ 列名全大写 + varchar(512)',
        'rows': '381,133',
        'cols': '26',
        'fields': [
            ['USRTY', 'varchar(512)', '类型', 'MAIL=邮箱 CELL=手机 MPHN=座机'],
            ['USRID', 'varchar(512)', '系统标识', '通讯号码'],
            ['USRID_LONG', 'varchar(512)', '长标识/号码', '完整通讯号码'],
        ]
    },
    {
        'name': 'ods_sap_pa0185',
        'desc': '人力资源主记录：信息类型 0185（证件信息）',
        'rows': '5,212',
        'cols': '43',
        'fields': [
            ['ictyp', 'varchar(2)', '证件类型', '01=身份证'],
            ['icnum', 'varchar(30)', '证件号码', ''],
            ['icold', 'varchar(20)', '旧证件号', ''],
            ['auth1', 'varchar(30)', '签发机构', ''],
            ['docn1', 'varchar(20)', '文件编号', ''],
            ['fpdat', 'varchar(8)', '指纹日期', ''],
            ['expid', 'varchar(8)', '有效期', ''],
            ['isspl', 'varchar(30)', '发证机关', ''],
            ['iscot', 'varchar(3)', '签发国', ''],
            ['idcot', 'varchar(3)', '证件国', ''],
            ['ovchk', 'varchar(1)', '验证标识', ''],
            ['astat', 'varchar(1)', '状态', ''],
            ['akind', 'varchar(1)', '证件种类', ''],
            ['rejec', 'varchar(20)', '拒绝原因', ''],
            ['usefr', 'varchar(8)', '有效期开始', 'yyyyMMdd'],
            ['useto', 'varchar(8)', '有效期截至', 'yyyyMMdd'],
            ['daten', 'decimal(3,0)', '天数', ''],
            ['dateu', 'varchar(3)', '日期单位', ''],
            ['times', 'varchar(8)', '时间', ''],
            ['hrguid', 'varchar(32)', '唯一标识', ''],
        ]
    },
    {
        'name': 'ods_sap_pa0534',
        'desc': '人力资源主记录：信息类型 0534（政治面貌）',
        'rows': '3,351',
        'cols': '40',
        'fields': [
            ['pcode', 'varchar(2)', '政治面貌代码', ''],
            ['joinu', 'varchar(40)', '加入组织', ''],
            ['intr1', 'varchar(8)', '介绍人1', ''],
            ['intr2', 'varchar(8)', '介绍人2', ''],
            ['paety', 'varchar(2)', '党员类型', ''],
            ['exprs', 'varchar(20)', '组织名称', ''],
            ['pstat', 'varchar(1)', '状态', ''],
            ['joipd', 'varchar(8)', '加入日期', ''],
            ['bfmdt', 'varchar(8)', '转正日期', ''],
            ['joitd', 'varchar(8)', '转入日期', ''],
            ['joimd', 'varchar(8)', '转出日期', ''],
            ['postx', 'varchar(40)', '岗位', ''],
            ['bdpos', 'varchar(8)', '开始岗位', ''],
            ['edpos', 'varchar(8)', '结束岗位', ''],
            ['intro', 'varchar(40)', '介绍人', ''],
            ['rema1', 'varchar(1000)', '备注1', ''],
            ['rema2', 'varchar(1000)', '备注2', ''],
        ]
    },
    {
        'name': 'ods_sap_pa105',
        'desc': '⚠ 冗余表 — 与 pa0105 结构相同，疑似 DataX 重复同步。小写列名 + varchar(n)',
        'rows': '15,897',
        'cols': '26',
        'fields': [
            ['usrty', 'varchar(4)', '类型', '同 PA0105.USRTY'],
            ['usrid', 'varchar(30)', '系统标识', '同 PA0105.USRID'],
            ['usrid_long', 'varchar(241)', '长标识/号码', '同 PA0105.USRID_LONG'],
        ]
    },
]

for tbl in tables_detail:
    doc.add_heading(f'{tbl["name"]} — {tbl["desc"]}', level=2)
    doc.add_paragraph(f'行数：{tbl["rows"]} | 业务列数：{len(tbl["fields"])} | 总列数：{tbl["cols"]}')

    table = doc.add_table(rows=len(tbl['fields'])+1, cols=4, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    for j, h in enumerate(['列名', '数据类型', '说明', '备注']):
        cell = table.rows[0].cells[j]
        cell.text = h
        for p in cell.paragraphs:
            p.runs[0].bold = True

    for i, field in enumerate(tbl['fields']):
        for j, val in enumerate(field):
            table.rows[i+1].cells[j].text = val

    doc.add_paragraph('')

# ============================================================
# 四、异常项
# ============================================================
doc.add_heading('四、异常项及建议', level=1)

anomalies = [
    ['异常', '涉及表', '说明', '建议'],
    ['空表 (0行)', 'ods_sap_pa0022\nods_sap_pa0024',
     'DataX 未同步到教育经历和资格认证数据。SAP 源端可能有数据但同步配置遗漏，或源端确实无数据。',
     '确认 SAP 源端是否有数据，如有则修复 DataX 同步任务'],
    ['列名大写 + varchar(512)', 'ods_sap_pa0105',
     '与其他 9 张表的命名风格不一致（其他为小写 + varchar(n) 精确类型）',
     '在 DWD 层清洗时统一为小写 + 精确类型'],
    ['疑似冗余表', 'ods_sap_pa105',
     '与 ods_sap_pa0105 结构相同但数据量不同（15,897 vs 381,133），'
     '可能是 DataX 不同批次的重复同步产物',
     '核实来源后删除冗余表或合并数据'],
    ['sources.yml 缺失', 'pa0002/pa0021/pa0022/\npa0023/pa0024/pa0185/\npa0534/pa105',
     '8 张表在 Hive 中存在但 dbt sources.yml 未定义，导致 dbt parse 编译失败',
     '补全 sources.yml 定义'],
]

table = doc.add_table(rows=len(anomalies), cols=4, style='Table Grid')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, row_data in enumerate(anomalies):
    for j, val in enumerate(row_data):
        cell = table.rows[i].cells[j]
        cell.text = val
        if i == 0:
            for p in cell.paragraphs:
                p.runs[0].bold = True

doc.add_paragraph('')

# ============================================================
# 五、Excel 需求 vs Hive 对照
# ============================================================
doc.add_heading('五、Excel 需求 vs Hive ODS 对照', level=1)

mapping = [
    ['Excel 模块', 'Excel 需求字段数', '对应 SAP 表', 'Hive 表名', '行数', '列数', '状态'],
    ['基本信息(组织)', '4 (姓名/公司/部门/职务)', 'PA0001', 'ods_sap_pa0001', '12,318', '57', '✅'],
    ['基本信息(个人)', '10 (性别/出生日期/民族等)', 'PA0002', 'ods_sap_pa0002', '5,286', '87', '✅'],
    ['基本信息(通讯)', '3 (邮箱/手机/座机)', 'PA0105', 'ods_sap_pa0105', '381,133', '26', '✅'],
    ['家庭成员', '19', 'PA0021', 'ods_sap_pa0021', '494', '80', '✅'],
    ['教育经历', '16', 'PA0022', 'ods_sap_pa0022', '0', '43', '⚠ 空表'],
    ['工作经历', '6', 'PA0023', 'ods_sap_pa0023', '2,799', '37', '✅'],
    ['资格认证', '7', 'PA0024', 'ods_sap_pa0024', '0', '25', '⚠ 空表'],
    ['证件信息', '6', 'PA0185', 'ods_sap_pa0185', '5,212', '43', '✅'],
    ['政治面貌', '1', 'PA0534', 'ods_sap_pa0534', '3,351', '40', '✅'],
]

table = doc.add_table(rows=len(mapping), cols=7, style='Table Grid')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, row_data in enumerate(mapping):
    for j, val in enumerate(row_data):
        cell = table.rows[i].cells[j]
        cell.text = val
        if i == 0:
            for p in cell.paragraphs:
                p.runs[0].bold = True

# ── 保存 ──
output_path = os.path.expanduser(r'~\Desktop\HR_ODS层表结构文档.docx')
doc.save(output_path)
print(f'已保存到: {output_path}')
