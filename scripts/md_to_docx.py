# -*- coding: utf-8 -*-
"""将 智能数据平台产品介绍.md 转换为精美 Word (.docx) 格式。
状态机解析器，支持嵌套列表、表格、代码块、链接、行内格式、截图占位。
"""

import re, os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

SRC = r"C:\Users\诗写\Desktop\dbt_warehouse\智能数据平台产品介绍.md"
DST = r"C:\Users\诗写\Desktop\dbt_warehouse\智能数据平台产品介绍.docx"

with open(SRC, "r", encoding="utf-8") as f:
    text = f.read()

doc = Document()

# ── 页面设置 ──
sec = doc.sections[0]
sec.page_width  = Cm(21)
sec.page_height = Cm(29.7)
sec.left_margin = Cm(2.2)
sec.right_margin = Cm(2.2)
sec.top_margin = Cm(2.0)
sec.bottom_margin = Cm(2.0)

# ── 颜色 ──
BLUE   = RGBColor(0x00, 0x7A, 0xFF)
GREEN  = RGBColor(0x34, 0xC7, 0x59)
ORANGE = RGBColor(0xFF, 0x95, 0x00)
RED    = RGBColor(0xFF, 0x3B, 0x30)
PURPLE = RGBColor(0xAF, 0x52, 0xDE)
TEAL   = RGBColor(0x5A, 0xC8, 0xFA)
GRAY   = RGBColor(0x86, 0x86, 0x8B)
DARK   = RGBColor(0x1D, 0x1D, 0x1F)
LGRAY  = RGBColor(0xAE, 0xAE, 0xB2)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
BG     = RGBColor(0xF5, 0xF5, 0xF7)

# ── 样式 ──
FONT = '微软雅黑'
MONO = 'Consolas'

for sn in ['Normal','Heading 1','Heading 2','Heading 3','Heading 4']:
    s = doc.styles[sn]
    s.font.name = FONT
    s.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)

ns = doc.styles['Normal']
ns.font.size = Pt(10.5)
ns.paragraph_format.space_after = Pt(6)
ns.paragraph_format.line_spacing = 1.5

for lv, sz, sp in [(1,22,12),(2,16,8),(3,13,6),(4,11,4)]:
    s = doc.styles[f'Heading {lv}']
    s.font.size = Pt(sz); s.font.bold = True
    s.paragraph_format.space_before = Pt(sp*2)
    s.paragraph_format.space_after = Pt(sp)
    if lv <= 2: s.font.color.rgb = DARK if lv==1 else BLUE

# ── 工具函数 ──

def run(p, text, color=None, bold=False, size=10.5, mono=False, italic=False):
    """段落追加带样式的 run"""
    r = p.add_run(text)
    if color: r.font.color.rgb = color
    if bold:  r.font.bold = True
    r.font.size = Pt(size)
    fn = MONO if mono else FONT
    r.font.name = fn
    r._element.rPr.rFonts.set(qn('w:eastAsia'), fn)
    if italic: r.font.italic = True
    return r

def new_para(size=10.5, indent=None, align=None, spacing_before=0, spacing_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(spacing_before)
    p.paragraph_format.space_after  = Pt(spacing_after)
    if indent: p.paragraph_format.left_indent = Cm(indent)
    if align:  p.alignment = align
    return p

def cell_shd(cell, hex_color):
    tcPr = cell._element.get_or_add_tcPr()
    tcPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>'))

def styled_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = 'Table Grid'
    for j,h in enumerate(headers):
        c = t.rows[0].cells[j]; c.text = ''
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run(p, h, WHITE, bold=True, size=9)
        cell_shd(c, "007AFF")
    for i,row in enumerate(rows):
        for j,val in enumerate(row):
            c = t.rows[i+1].cells[j]; c.text = ''
            p = c.paragraphs[0]
            run(p, str(val), size=9)
            if i % 2 == 0: cell_shd(c, "F5F5F7")
    if col_widths:
        for row in t.rows:
            for j,w in enumerate(col_widths):
                if j < len(row.cells): row.cells[j].width = Cm(w)
    doc.add_paragraph()
    return t

def code_block(lines_list):
    for cl in lines_list:
        p = new_para(size=9, indent=0.5, spacing_before=2, spacing_after=2)
        pPr = p._element.get_or_add_pPr()
        pPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F7" w:val="clear"/>'))
        run(p, cl, DARK, size=9, mono=True)

def hl_para(raw, size=10.5, indent=None, bold_all=False):
    """行内格式解析: **bold** `code` [text](url)  → 追加到段落"""
    p = new_para(size=size, indent=indent)
    # tokenize
    tokens = re.split(r'(\*\*[^*]+\*\*|`[^`]+`|\[([^\]]+)\]\(([^)]+)\))', raw)
    i = 0
    while i < len(tokens):
        t = tokens[i]
        if t is None: i+=1; continue
        if t.startswith('**') and t.endswith('**'):
            run(p, t[2:-2], bold=True, size=size)
        elif t.startswith('`') and t.endswith('`'):
            run(p, t[1:-1], size=size-1, mono=True)
        elif i+2 < len(tokens) and tokens[i+1] is not None and tokens[i+2] is not None:
            # link: text=token[i+1], url=token[i+2]
            lnk_text = tokens[i+1]
            lnk_url  = tokens[i+2]
            if lnk_text:
                run(p, lnk_text, BLUE, size=size)
            i += 2
        elif t.strip():
            run(p, t, size=size, bold=bold_all)
        i += 1
    return p

def hr():
    p = new_para(spacing_before=10, spacing_after=10)
    pPr = p._element.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="D9D9D9"/></w:pBdr>'))

def img_placeholder(alt, src):
    p = new_para(align=WD_ALIGN_PARAGRAPH.CENTER, spacing_before=8)
    run(p, f'📷 [截图占位: {alt}]', GRAY, size=10, italic=True)
    p2 = new_para(size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    run(p2, src, LGRAY, size=8)

def quote(text):
    p = new_para(size=9, indent=1.0, spacing_before=2, spacing_after=2)
    pPr = p._element.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:pBdr {nsdecls("w")}><w:left w:val="single" w:sz="12" w:space="8" w:color="007AFF"/></w:pBdr>'))
    tokens = re.split(r'(\*\*[^*]+\*\*)', text)
    for t in tokens:
        if t.startswith('**') and t.endswith('**'):
            run(p, t[2:-2], bold=True, size=9, color=GRAY)
        else:
            run(p, t, GRAY, size=9)

# ── 状态机主解析 ──
lines = text.split('\n')
i = 0
in_code = False; code_buf = []
table_buf = []    # list of (headers, rows) waiting to render after current section
pending_hdrs = None; pending_rows = []

while i < len(lines):
    line = lines[i]
    stripped = line.strip()

    # 空行
    if not stripped:
        if in_code: code_buf.append('')
        # flush pending table if next non-empty line is NOT a table continuation
        if pending_hdrs and i+1 < len(lines):
            nxt = lines[i+1].strip()
            if not (nxt.startswith('|') and nxt.endswith('|')):
                styled_table(pending_hdrs, pending_rows)
                pending_hdrs = None; pending_rows = []
        i += 1; continue

    # 代码块
    if stripped.startswith('```'):
        if in_code:
            code_block(code_buf)
            code_buf = []; in_code = False
        else:
            in_code = True
        i += 1; continue

    if in_code:
        code_buf.append(line)
        i += 1; continue

    # 水平线
    if stripped == '---':
        # flush any pending table
        if pending_hdrs:
            styled_table(pending_hdrs, pending_rows)
            pending_hdrs = None; pending_rows = []
        hr()
        i += 1; continue

    # 标题
    m = re.match(r'^(#{1,4})\s+(.+)', stripped)
    if m:
        if pending_hdrs:
            styled_table(pending_hdrs, pending_rows)
            pending_hdrs = None; pending_rows = []
        lv = len(m.group(1))
        doc.add_heading(m.group(2).strip(), level=lv)
        i += 1; continue

    # 引用块
    if stripped.startswith('> '):
        if pending_hdrs:
            styled_table(pending_hdrs, pending_rows)
            pending_hdrs = None; pending_rows = []
        quote(stripped[2:])
        i += 1; continue

    # 图片
    if stripped.startswith('!['):
        alt = stripped[2:stripped.index(']')]
        src = stripped[stripped.index('(')+1:stripped.index(')')]
        img_placeholder(alt, src)
        i += 1; continue

    # 表格行
    if stripped.startswith('|') and stripped.endswith('|'):
        cells = [c.strip() for c in stripped.split('|')[1:-1]]
        # 分隔行 `| --- | --- |`
        if all(re.match(r'^:?-{3,}:?$', c) for c in cells):
            i += 1; continue
        if pending_hdrs is None:
            pending_hdrs = cells
            pending_rows = []
        else:
            pending_rows.append(cells)
        # peek next line
        if i+1 >= len(lines) or not (lines[i+1].strip().startswith('|') and lines[i+1].strip().endswith('|')):
            styled_table(pending_hdrs, pending_rows)
            pending_hdrs = None; pending_rows = []
        i += 1; continue

    # 多级无序列表
    m_list = re.match(r'^(\s*)[-*]\s+(.+)', line)
    if m_list:
        if pending_hdrs:
            styled_table(pending_hdrs, pending_rows)
            pending_hdrs = None; pending_rows = []
        indent_lv = len(m_list.group(1))
        content = m_list.group(2)
        p = doc.add_paragraph(style='List Bullet')
        p.clear()
        p.paragraph_format.left_indent = Cm(1.2 + indent_lv * 0.6)
        # inline parse
        tokens = re.split(r'(\*\*[^*]+\*\*|`[^`]+`|\[([^\]]+)\]\(([^)]+)\))', content)
        j = 0
        while j < len(tokens):
            t = tokens[j]
            if t is None: j+=1; continue
            if t.startswith('**') and t.endswith('**'):
                run(p, t[2:-2], bold=True, size=10.5)
            elif t.startswith('`') and t.endswith('`'):
                run(p, t[1:-1], size=9, mono=True)
            elif j+2 < len(tokens) and tokens[j+1] is not None and tokens[j+2] is not None:
                run(p, tokens[j+1], BLUE, size=10.5)
                j += 2
            elif t.strip():
                run(p, t, size=10.5)
            j += 1
        # 处理多行列表项（后续行以空格开头且不是新列表项）
        while i+1 < len(lines) and lines[i+1].startswith('  ') and not re.match(r'^\s*[-*]\s', lines[i+1]):
            i += 1
            cont = lines[i].strip()
            if cont:
                p2 = new_para(size=10.5, indent=2.0)
                run(p2, cont, size=10.5)
        i += 1; continue

    # 有序列表
    m_ol = re.match(r'^(\s*)\d+\.\s+(.+)', line)
    if m_ol:
        if pending_hdrs:
            styled_table(pending_hdrs, pending_rows)
            pending_hdrs = None; pending_rows = []
        content = m_ol.group(2)
        p = doc.add_paragraph(style='List Number')
        p.clear()
        tokens = re.split(r'(\*\*[^*]+\*\*|`[^`]+`|\[([^\]]+)\]\(([^)]+)\))', content)
        j = 0
        while j < len(tokens):
            t = tokens[j]
            if t is None: j+=1; continue
            if t.startswith('**') and t.endswith('**'):
                run(p, t[2:-2], bold=True, size=10.5)
            elif t.startswith('`') and t.endswith('`'):
                run(p, t[1:-1], size=9, mono=True)
            elif j+2 < len(tokens) and tokens[j+1] is not None and tokens[j+2] is not None:
                run(p, tokens[j+1], BLUE, size=10.5)
                j += 2
            elif t.strip():
                run(p, t, size=10.5)
            j += 1
        i += 1; continue

    # 普通段落
    if pending_hdrs:
        styled_table(pending_hdrs, pending_rows)
        pending_hdrs = None; pending_rows = []
    hl_para(stripped)
    i += 1

# flush remaining table
if pending_hdrs:
    styled_table(pending_hdrs, pending_rows)

doc.save(DST)
print(f"OK — {DST}  ({os.path.getsize(DST)/1024:.1f} KB)")
